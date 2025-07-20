# FILE: app/services/document_processor.py

import logging
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
import re

from app.models.rag_models import DocumentChunk, ProcessedDocument, DocumentProcessingStatus
from app.core.firebase import db
from app.services.vertex_rag_service import VertexAIRAGService

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing documents and creating text chunks."""
    
    def __init__(self):
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
        self.processed_docs_collection = "processed_documents"
        self.document_chunks_collection = "document_chunks"
        self.vector_service = VertexAIRAGService()
    
    async def process_document(
        self,
        file_path: str,
        document_id: str,
        teacher_uid: str,
        filename: str,
        subject: Optional[str] = None,
        grade_level: Optional[int] = None
    ) -> ProcessedDocument:
        """Process a document and create chunks."""
        
        try:
            logger.info(f"Starting document processing: {document_id}")
            
            # Create initial document record
            processed_doc = ProcessedDocument(
                document_id=document_id,
                teacher_uid=teacher_uid,
                original_filename=filename,
                file_type=self._get_file_type(filename),
                subject=subject,
                grade_level=grade_level,
                processing_status=DocumentProcessingStatus.PROCESSING
            )
            
            # Save initial status
            await self._save_processed_document(processed_doc)
            
            # Extract text based on file type
            text_content = await self._extract_text(file_path, filename)
            
            if not text_content.strip():
                raise ValueError("No text content found in document")
            
            # Create chunks
            chunks = await self._create_chunks(
                text_content=text_content,
                document_id=document_id,
                metadata={
                    "subject": subject,
                    "grade_level": grade_level,
                    "teacher_uid": teacher_uid,
                    "filename": filename
                }
            )
            
            # Save chunks to database and vector store
            await self._save_chunks_to_vector_store(chunks)
            
            # Update document status
            processed_doc.total_chunks = len(chunks)
            processed_doc.processing_status = DocumentProcessingStatus.COMPLETED
            processed_doc.processed_at = processed_doc.created_at
            
            await self._save_processed_document(processed_doc)
            
            logger.info(f"Document processing completed: {document_id}, chunks: {len(chunks)}")
            return processed_doc
            
        except Exception as e:
            logger.error(f"Document processing failed for {document_id}: {str(e)}")
            
            # Update document with error status
            processed_doc.processing_status = DocumentProcessingStatus.FAILED
            processed_doc.processing_error = str(e)
            await self._save_processed_document(processed_doc)
            
            raise e
    
    async def _extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from various file types."""
        
        file_type = self._get_file_type(filename)
        
        try:
            if file_type == "pdf":
                return await self._extract_pdf_text(file_path)
            elif file_type == "docx":
                return await self._extract_docx_text(file_path)
            elif file_type == "txt":
                return await self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            raise e
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text_content = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\\n[Page {page_num + 1}]\\n{page_text}\\n"
                        
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise e
            
        return text_content.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        text_content = ""
        
        try:
            doc = DocxDocument(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\\n"
                    
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise e
            
        return text_content.strip()
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise e
    
    async def _create_chunks(
        self,
        text_content: str,
        document_id: str,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create overlapping text chunks from document content."""
        
        chunks = []
        
        # Clean and normalize text
        cleaned_text = self._clean_text(text_content)
        
        # Split into sentences first
        sentences = self._split_into_sentences(cleaned_text)
        
        current_chunk = ""
        current_chunk_sentences = []
        chunk_index = 0
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
                current_chunk_sentences.append(sentence)
            else:
                # Save current chunk if it has content
                if current_chunk.strip():
                    chunk = DocumentChunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=document_id,
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        metadata=metadata
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_sentences = current_chunk_sentences[-self._calculate_overlap_sentences():]
                current_chunk = " ".join(overlap_sentences + [sentence])
                current_chunk_sentences = overlap_sentences + [sentence]
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                chunk_id=str(uuid.uuid4()),
                document_id=document_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                metadata=metadata
            )
            chunks.append(chunk)
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = re.sub(r'\\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-\"\']', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting (could be enhanced with NLTK)
        sentence_endings = r'[.!?]+(?=\\s+[A-Z]|$)'
        sentences = re.split(sentence_endings, text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Filter out very short fragments
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _calculate_overlap_sentences(self) -> int:
        """Calculate number of sentences to include in overlap."""
        # Rough estimate: if chunk_overlap is 200 chars and avg sentence is ~50 chars
        return max(1, self.chunk_overlap // 50)
    
    def _get_file_type(self, filename: str) -> str:
        """Get file type from filename."""
        return Path(filename).suffix.lower().lstrip('.')
    
    async def _save_processed_document(self, doc: ProcessedDocument) -> None:
        """Save processed document metadata to Firestore."""
        try:
            doc_ref = db.collection(self.processed_docs_collection).document(doc.document_id)
            doc_ref.set(doc.dict())
            logger.info(f"Saved processed document: {doc.document_id}")
        except Exception as e:
            logger.error(f"Failed to save processed document {doc.document_id}: {str(e)}")
            raise e
    
    async def _save_chunks(self, chunks: List[DocumentChunk]) -> None:
        """Save document chunks to Firestore."""
        try:
            batch = db.batch()
            
            for chunk in chunks:
                doc_ref = db.collection(self.document_chunks_collection).document(chunk.chunk_id)
                batch.set(doc_ref, chunk.dict())
            
            batch.commit()
            logger.info(f"Saved {len(chunks)} chunks to Firestore")
            
        except Exception as e:
            logger.error(f"Failed to save chunks: {str(e)}")
            raise e
    
    async def _save_chunks_to_vector_store(self, chunks: List[DocumentChunk]) -> None:
        """Save document chunks to vector store with embeddings."""
        try:
            # Use the vector service to add chunks with embeddings
            success = await self.vector_service.add_chunks(chunks)
            
            if success:
                logger.info(f"Saved {len(chunks)} chunks to vector store with embeddings")
            else:
                logger.error("Failed to save chunks to vector store")
                raise Exception("Vector store save failed")
                
        except Exception as e:
            logger.error(f"Failed to save chunks to vector store: {str(e)}")
            raise e
    
    async def get_processed_document(self, document_id: str) -> Optional[ProcessedDocument]:
        """Get processed document metadata."""
        try:
            doc_ref = db.collection(self.processed_docs_collection).document(document_id)
            doc = doc_ref.get()
            
            if doc.exists:
                return ProcessedDocument(**doc.to_dict())
            return None
            
        except Exception as e:
            logger.error(f"Failed to get processed document {document_id}: {str(e)}")
            return None
    
    async def get_teacher_documents(self, teacher_uid: str) -> List[ProcessedDocument]:
        """Get all processed documents for a teacher."""
        try:
            docs_ref = db.collection(self.processed_docs_collection)
            query = docs_ref.where("teacher_uid", "==", teacher_uid).order_by("created_at", direction="DESCENDING")
            
            docs = query.stream()
            
            processed_docs = []
            for doc in docs:
                processed_docs.append(ProcessedDocument(**doc.to_dict()))
                
            return processed_docs
            
        except Exception as e:
            logger.error(f"Failed to get teacher documents for {teacher_uid}: {str(e)}")
            return []
